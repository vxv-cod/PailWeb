from fileinput import filename
from smtplib import SMTP_SSL
from PyQt5 import QtCore, QtGui, QtWidgets
import os
import sys
import pickle
# from okno_general import ui
# from okno_general import Form
# from okno_general import app

# from Tab73 import fsi73

# import docx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
document = Document()

# Рисунок
from PyQt5.QtWidgets import QWidget
from PyQt5.QtWidgets import QApplication
from PyQt5.QtGui import QPainter
from PyQt5.QtGui import QBrush
from PyQt5.QtGui import QPen
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap

# import vxv_tnnc_SQL_dll

# _translate = QtCore.QCoreApplication.translate
# Form.setWindowTitle(_translate("Form", "Pail"))
# ui.tabWidget.setCurrentIndex(0)

# -----------------------------------------------------------------------------
def insertstylcss(x):
    # with open("777777777777.html", "a", encoding='utf-8') as file:
    #     file.write(x)
    #     file.close()    
    global textOtchet
    textOtchet += x




            
def text_centr(x):
    xxx = f'''
        <p style="margin-top:0pt; margin-bottom:10pt; text-align:center; font-weight:bold; font-style:italic; color:#4f81bd">
        {x}</p>
    '''
    insertstylcss(xxx)
    # ui.textEdit.append('')
    # ui.textEdit.setTextColor(QtGui.QColor (0, 100, 150)) # цвет текста
    # # ui.textEdit.setFontItalic(True) # курсивный текст
    # ui.textEdit.setFontWeight(100) # жирный текст
    # ui.textEdit.append(x)
    # ui.textEdit.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    # ui.textEdit.setFontWeight(1) # убираем жирный текст
    # # ui.textEdit.setFontItalic(False) # убираем курсивный текст
    # ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
    global document
    # document.add_paragraph(x)
    paragraph = document.add_paragraph()
    # paragraph.add_run(x).bold = True
    paragraph.add_run(x, style='Intense Emphasis').bold = True
    # paragraph.add_run(x).font.color.rgb = RGBColor(0, 100, 150)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

def text_centr_black(x):
    xxx = f'''
    <p style="margin-top:0pt; margin-bottom:10pt; text-align:center">
    <span>{x}</span>
    </p>
    '''
    insertstylcss(xxx)
    
    # ui.textEdit.append(x)
    # ui.textEdit.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    global document
    paragraph = document.add_paragraph(x)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def text_abzac(x):
    # xxx = f'<p name:"text_abzac"; style="margin-top:0pt; margin-bottom:10pt; width:36pt; display:inline-block">{x}</p>'
    xxx = f'''
    <p style="margin-top:0pt; margin-bottom:10pt">
        <span style="width:36pt; display:inline-block">&#xa0;</span>
        <span>{x}</span></p>
    '''
    insertstylcss(xxx)
    
    # ui.textEdit.append('       {}'.format(x))
    # ui.textEdit.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    global document
    document.add_paragraph(x)
    
def text_abzac_color(x):
    xxx = f'''
    <p style="margin-top:0pt; margin-bottom:10pt">
        <span style="font-weight:bold; font-style:italic; color:#4f81bd">{x}</span>
    </p>
    '''
    insertstylcss(xxx)
    
    # ui.textEdit.append('')
    # ui.textEdit.setTextColor(QtGui.QColor (0, 100, 150))
    # ui.textEdit.setFontWeight(100) # жирный текст
    # ui.textEdit.append('       {}'.format(x))
    # ui.textEdit.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    # ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
    # ui.textEdit.setFontWeight(1) # жирный текст
    # ui.textEdit.append('')
    global document
    paragraph = document.add_paragraph()
    paragraph.add_run(x, style='Intense Emphasis').bold = True

# -----------------------------------------------------------------------------
# # Проверка ячеек на пустые значения и ","
# def vvod(nomerwidgeta, strok, stolbec):
#     xox = []; yoy = []
#     for i in range(1, strok+1):
#         x = eval('ui.tableWidget{}.item({}-1, {}).text()'.format(nomerwidgeta, i, stolbec))
#         if x != '': 
#             x = x.replace(',', '.')
#         else: 
#             x = 0
#         xox.append(x)
#         if '.' in str(x):
#             yoy.append(float(x))
#         else:
#             try:
#                 yoy.append(int(x))
#             except:
#                 yoy.append(str(x))
#     return xox, yoy

# def sbor_dannih():
#     global ige_xap, qi_xap, fi_xap, dannie, ige_skv, nni, sechenie
#     sechenie = ui.radioButton.isChecked()
#     # Ввод параметров грунта
#     ige_xap = vvod('', 20, 0)[0]
#     qi_xap = vvod('', 20, 1)[1]
#     fi_xap = vvod('', 20, 2)[1]
#     dannie = vvod('_2', 19, 0)[1]       # исходные данные
#     ige_skv = vvod('_4', 18, 0)[0]      # список ИГЭ скважины
#     nni = vvod('_4', 18, 1)[1]          # список толщин слоев скважины

# # Вывод ошибки
# def error_show(x):
# def print(x):
#     ui.label_3.setText(_translate("Form", "Ошибка"))
#     ui.textEdit.setTextColor(QtGui.QColor (255, 0, 0))
#     ui.textEdit.setText('Ошибка данных:')
#     ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
#     ui.textEdit.append(x)
     
# -----------------------------------------------------------------------------
def raschet(dannie, ige_skv, nni, ige_xap, qi_xap, fi_xap):
    '''Отправляем сигнал на сервер об нажатии кнопки "Расчет" '''
    # vxv_tnnc_SQL_dll.SendText('Pail_Stadzond')

    # _translate = QtCore.QCoreApplication.translate
    # ui.label_3.setText(_translate("Form", "Отчет"))
    global textOtchet
    global document
    del document    # удаляем глобальную переменную и заменяем ее новой такой же, но пустой
    document = Document()

    # with open("777777777777.html", "w+", encoding='utf-8') as file:
    #     file.write("text")
    #     file.close()
        
    textOtchet = ''
    
    # sbor_dannih()
    # global ige_xap, qi_xap, fi_xap, dannie, ige_skv, nni, sechenie
    # global ige_xap, qi_xap, fi_xap, dannie, ige_skv, nni    
    # global nni
    print(f"nni_raschet = {nni}")
    
    
    
    
    
    
    
    
    
    
    nni_ishod = nni
    for i in range(0, len(ige_skv)):
        if ige_skv[i] not in ige_xap:
            # error_show('ИГЭ скважины не соответствуют ИГЭ в таблице \" Характеристики грунтов \"')
            print('ИГЭ скважины не соответствуют ИГЭ в таблице \" Характеристики грунтов \"')
            return
    qi_dict = {}
    for i in range(0, len(ige_xap)):
        qi_dict[ige_xap[i]] = qi_xap[i]
    fi_dict = {}
    for i in range(0, len(ige_xap)):
        fi_dict[ige_xap[i]] = fi_xap[i]

    qi = [qi_dict[i] for i in ige_skv]    # список сопротивления слоев грунта под конусом
    fi = [fi_dict[i] for i in ige_skv]    # список сопротивления слоев грунта по муфте   
    psi = [2 if qi[i] >= dannie[16] else 1 for i in range (0, len(qi))]      # список назначения песчаных или глинистых грунтов
    torf = [1 if 0 < qi[i] <= dannie[15] else 0 for i in range (0, len(qi))]      # список торфяных грунтов
    torf_ishod = torf
    ace = dannie[14] # глубина промерзания
    kru = dannie[9] # понижающий коэффициент антисмерзающего покрытия сваи

    #===============================================================================
    # Ввод параметров сваи
    d1 = dannie[2]*0.001  # Внешний диаметр сваи
    tsv = dannie[3]*0.001 # Толщина стенки сваи
    l1 = dannie[1]        # Длина сваи
    l1_ishod = l1
    l2 = float(dannie[0])          # Отметка сваи
    l2_ishod = l2
    l3 = l1 - l2            # Отметка низа сваи
    l3_ishod = l3
    l4 = round(sum(nni) - l3, 3)                          # Под нижним концом
    # _translate = QtCore.QCoreApplication.translate
    # ui.tableWidget_2.item(18, 0).setText(_translate("Form", str(l4)))

    if sum(nni) == 0:
        # error_show('Не заполнены данные!!!')
        print('Не заполнены данные!!!')
        return
    ni_geo = nni[:]
    while ni_geo[-1] == 0: del ni_geo[-1]

    app = QtWidgets.QApplication(sys.argv)


    # Прорисовка Геологического разреза по данным 
    class Example(QWidget):
        def __init__(self, nni, brs, l1, l2, Hrisunka):
            super().__init__()
            self.initUI()
        def initUI(self):
            self.setGeometry(0, 0, 180, Hrisunka)
            self.setStyleSheet("background-color: rgb(0, 0, 0, 0);")
            # self.setWindowTitle('Brushes')
            # self.show()
        def paintEvent(self, e):
            qp = QPainter()
            qp.begin(self)
            self.drawBrushes(qp, nni, brs)
            qp.end()
        def drawBrushes(self, qp, nni, brs):
            fgh = l2 * 30
            a = fgh + 15 if fgh >= 0 else 15
            sv0 = 15 if fgh > 0 else 15 - fgh
            sv1 =  l1 * 30 + sv0
            qp.setPen(QPen(QtGui.QColor(0,0,0), 5, QtCore.Qt.SolidLine))
            qp.drawLine(90, sv0, 90, sv1)
            brush = QBrush(Qt.NoBrush)
            brush.setColor(QtGui.QColor(0,100,150,150))
            qp.setPen(QPen(QtGui.QColor(0,0,0), 1, QtCore.Qt.SolidLine))

            qp.drawText(135, a+5, '0.000')
            qp.drawText(0, sv0-6, 45, 12, Qt.AlignRight, ('+ ' if fgh >= 0 else '- ') + str(round(abs(l2), 3)))
            qp.drawText(0, sv1-6, 45, 12, Qt.AlignRight, '- ' + str(round(abs(l3), 3)))

            qp.setPen(QPen(QtGui.QColor(0,0,0), 1, QtCore.Qt.DashLine))
            qp.drawLine(50, sv0, 80, sv0)
            qp.drawLine(50, sv1, 80, sv1)
            qp.setPen(QPen(QtGui.QColor(0,0,0), 1, QtCore.Qt.SolidLine))
            
            for i in range(0, len(ni_geo)):
                b = float(nni[i])*30
                brush.setStyle(eval(brs[i]))
                qp.setBrush(brush)
                qp.drawRect(55, a, 70, b)
                a += b
                qp.drawText(135, a+5, '- ' + str(round((a-fgh-15)/30, 3)) if fgh >= 0 else '- ' + str(round((a-15)/30, 3)))
    
    tigr = []
    for i in range(0, len(psi)):
        if psi[i] == 1 and torf_ishod[i] == 0:
            tigr.append('глинистый')
        if psi[i] == 1 and torf_ishod[i] == 1:
            tigr.append('торф')
        if psi[i] == 2 and torf_ishod[i] == 0 and fi_xap[i] != 0:
            tigr.append('песчаный')
        if psi[i] == 2 and torf_ishod[i] == 0 and fi_xap[i] == 0:
            tigr.append('насыпь')

    # tigr =  ['насыпь', 'торф', 'глинистый', 'песчаный']
    brs = []
    for i in range(0, len(tigr)):
        if tigr[i] == 'насыпь': brs.append('Qt.DiagCrossPattern')
        if tigr[i] == 'торф': brs.append('Qt.CrossPattern')
        if tigr[i] == 'глинистый': brs.append('Qt.BDiagPattern')
        if tigr[i] == 'песчаный': brs.append('Qt.Dense5Pattern')

    Hrisunka = sum(nni)*30 + 30 + (l2 * 30 if l2 >= 0 else 0)
    ex = Example(nni, brs, l1, l2, Hrisunka)
    # ex.setStyleSheet("background-color: rgb(254, 254, 254);")
    ee = ex.grab()
    pixmap = QPixmap(ee)
    # ui.label_10.setGeometry(QtCore.QRect(0, 0, 190, Hrisunka))
    # ui.label_10.setPixmap(pixmap)

    # ex.setStyleSheet("background-color: rgb(255, 255, 255);")
    # ex.setStyleSheet("background-color: rgb(181, 180, 180, 100);")
    ex.setStyleSheet("background-color: rgb(255, 255, 255, 0);")
    ee = ex.grab()
    pix = QPixmap(ee)

    # return

    # Вывод ошибки при не корректных исходных данных:
    if l3 >= sum(nni):
        # error_show('Глубина погружения сваи в грунт равна или превышает тошщину всех слоев грунта')
        print('Глубина погружения сваи в грунт равна или превышает тошщину всех слоев грунта')
        return

    if l3 <= l2:
        # error_show('Длина сваи недостаточна!!!')
        print('Длина сваи недостаточна!!!')
        return

    #===============================================================================
    Ns = dannie[6]  # Введите сжимающую нагрузку на сваю, тс
    Nv = dannie[7]  # Введите выдергивающую нагрузку на сваю, тс
    zsv = dannie[4] # Введите объемный вес заполнения сваи, тс/м3
    knn_met = dannie[11] # Коэффициент надежности по нагрузки для металла
    knn_beton = dannie[12] # Коэффициент надежности по нагрузки для бетона
    #===============================================================================
    # return
    # Расчетные толщины слоев грунта около поверхности сваи
    def funni (nni, l3):
        ni = nni [:]
        while sum (ni) > l3: del ni [-1]
        ni.append (l3 - sum (ni)); return ni
    ni = funni (nni, l3) # обрезаем грунт под сваей
    ni_ishod = ni
    ni_tolchina = ni
    ni_tolchina = funni (nni_ishod, l3_ishod)

    def verx_niz(ni, l2):
        l3 = abs (l2); verx = funni (nni, l3) # глубина промерзания по свае
        while len(verx) < len(ni): verx.append (0)
        niz = [ni[i] - verx[i] for i in range(0, len(ni))]
        return verx, niz

    if l2 < 0:
        ni = verx_niz(ni, l2)[1] # глубина без промерзания по свае
        ace = 0 if ace + l2 <= 0 else ace + l2
        l2 = ace; nni = ni[:]
    else: l2 = ace
    v_n = verx_niz(ni, l2)
    ni = v_n[1] # грунт по свае без промерзания
    ni_tddd = ni

    if sum(ni) < 0:
        # error_show('Длина сваи недостаточна!!!')
        print('Длина сваи недостаточна!!!')
        return

    niend = round(ni[-1], 3)
    # ui.tableWidget_2.item(17, 0).setText(_translate("Form", str(niend)))    # Кусок сваи в последнем слое

    # Модуль расчета 1
    # if ui.radioButton.isChecked() == True:
    if True:
        asv = 3.14159265359 * d1*d1 * 0.25                  # Площадь поперечного сечения сваи, м2
        usv = 3.14159265359 * d1                            # Периметр сечения сваи, м
        R = d1*0.5                                          # Внешний радиус сваи
        r = (d1 - 2*tsv)*0.5                                # Внутренний радиус сваи
        msv = 3.14159265359 * (R**2 - r**2) * l1 * 7.85     # Вес металла сваи
        bsv = 3.14159265359 * r**2 *  l1 * zsv              # Вес заполняющего полость бетона
        vsv = msv * knn_met + bsv * knn_beton
    else:
        asv = d1 * d1           # Площадь поперечного сечения сваи, м2
        usv = d1 * 4            # Периметр сечения сваи, м
        bsv = asv *  l1 * zsv
        vsv = bsv * knn_beton   # Вес бетонной сваи

    kk = len (ni); fi = fi [:kk]; qi = qi [:kk]; psi = psi [:kk]; torf = torf [:kk]     # Корректировка спиков по длине спика ni
    fi = [i * 0.1019716212978 for i in fi]      # Перевод ед. изм. с кПа на тс/м2 в списках
    qs = qi[-1] * 101.9716212978                # Среднее значение сопротивления грунта под наконечником зонда
    torf = torf [:kk]                           # Корректировка спиков по длине спика ni
    if sum(torf) > 0:
        while torf[-1] == 0:
            del torf[-1]
        ni_osadki = [nni[i] for i in range(0, len(torf))]
        ni = ni_ishod

        for i in range(len(ni)):
            if i <= len(ni_osadki)-1:
                ni[i] = 0

    def osadka_torfa(bi_tab, ni_osadki, fi, psi, torf, usv):
        fi_osadki = [fi[i] for i in range(0, len(torf))]
        psi_osadki = [psi[i] for i in range(0, len(torf))]
        ni_torfa = [torf[i] * ni_osadki[i] for i in range(0, len(torf))]
        ni_peska = [ni_osadki[i] - ni_torfa[i] for i in range(0, len(torf))]
        fi_torfa = [torf[i] * fi_osadki[i] for i in range(0, len(torf))]
        fi_peska = [fi_osadki[i] - fi_torfa[i] for i in range(0, len(torf))]
        Htorfa = sum(ni_torfa)
        Hpeska = sum(ni_peska)
                
        if Htorfa >= 0.3 and Hpeska >= 2.0:
            for i in range(0, len(ni_osadki)):
                if ni_peska[i] >= 2.0 and fi_peska[i] == 0:
                    xx =[2.0 for x in range(0, int(ni_osadki[i]//2))]
                    if ni_osadki[i] % 2 != 0.0:
                        xx.append(ni_osadki[i] % 2)
                    ni_osadki[i] = xx

            fi_73 = []; jp = 0.3; l = 0
            bi_osadki = []
            for i in range(0, len(ni_osadki)):
                if type(ni_osadki[i]) == list:
                    fi_73_i = []
                    for x in range(0, len(ni_osadki[i])):
                        if fi_osadki[i] == 0 and psi_osadki[i] == 2:
                            l = l + ni_osadki[i][x] * 0.5
                            if 2.0 <= Hpeska < 5.0:
                                fi_73_i.append(fsi73(l, jp)*0.4)
                            else: fi_73_i.append(fsi73(l, jp))
                            l = l + ni_osadki[i][x] * 0.5
                    bi_osadki.append(1.0)
                    fi_73.append(fi_73_i)
                else:
                    l = l + ni_osadki[i] * 0.5
                    if fi_osadki[i] == 0 and psi_osadki[i] == 2:
                        fi_73.append(fsi73(l, jp))
                        bi_osadki.append(1.0)
                    if fi_osadki[i] != 0 and psi_osadki[i] == 2:
                        fi_73.append(fi_osadki[i])
                        bi_osadki.append(bi_tab[i])
                    if fi_osadki[i] == 0 and psi_osadki[i] == 1:
                        fi_73.append(0.509858)
                        bi_osadki.append(1.0)
                    if fi_osadki[i] != 0 and psi_osadki[i] == 1:
                        fi_73.append(fi_osadki[i])
                        bi_osadki.append(1.0)
                    l = l + ni_osadki[i] * 0.5
            fi_osadki = fi_73
        
            bfh = []; sum_bfh_i = []
            for i in range(0, len(ni_osadki)):
                if type(ni_osadki[i]) == list:
                    bfh_i_i = []
                    for x in range(0, len(ni_osadki[i])):
                        bfh_i = round(bi_osadki[i] * fi_osadki[i][x] * ni_osadki[i][x], 3)
                        bfh_i_i.append(bfh_i)
                        sum_bfh_i.append(bfh_i)
                    bfh.append(bfh_i_i)
                else: 
                    bfh_i = round(bi_osadki[i] * fi_osadki[i] * ni_osadki[i], 3)
                    bfh.append(bfh_i)
                    sum_bfh_i.append(bfh_i)

            Fosadki = sum(sum_bfh_i) * usv

            return Fosadki, bi_osadki, fi_osadki, ni_osadki, bfh_i, bfh, sum_bfh_i
        else: 
            Fosadki = 0
            return Fosadki
    #===============================================================================
    # Функция интерполяции 
    def interpoi (t1, t2, yy):
        for i in t1:
            if yy > i:
                continue
            else:
                break
        ia1 = t1.index(i); ia2 = t1.index(i)-1
        a1 = t1[ia1]; a2 = t1[ia2]; b1 = t2[ia1]; b2 = t2[ia2]
        a12 = a1 - a2; b12 = b1 - b2; a13 = yy - a2; x = a13 * b12 / a12; y = b2 + x
        return y
    #===============================================================================
    # Коэффициент перехода от fsi к f для зонда типа II или III, βi по таблице 7.16 СП 24.13330-2011 изм.1
    def Tab716 (x, yy):
        t1 = 2.039432425956, 4.078864851912, 6.118297277868, 8.157729703824, 10.19716212978, 12.236594555736
        t21 = 1.0, 0.75, 0.6, 0.45, 0.4, 0.3 # при глинистых грунтах
        t22 = 0.75, 0.6, 0.55, 0.5, 0.45, 0.4 # при песчаных грунтах
        dct = {1:t21, 2:t22}
        t2 = dct[x]
        if yy < 2.039432425956:
            yy = 2.039432425956
        return interpoi (t1, t2, yy)
    #===============================================================================
    #Суммарное значение сопротивления грунта на боковой поверхности сваи f
    def bii(fi):
        bi = []
        for yy in fi:
            x = psi[fi.index(yy)]
            i = Tab716 (x, yy); bi.append (i)
        return bi
    def fsf (bi, ni, fi):
        # bi = bii(fi)
        f = []
        for i in range (0, len(ni)):
            f.append (bi[i]*ni[i]*fi[i])
        f = sum (f)
        return f
    #===============================================================================
    # Нормативное значение предельного сопротивления сваи на выдергивание, тс 
    def fu_n ():
        bi = bii(fi)
        return fsf (bi, ni, fi) * usv
    Funv = fu_n ()
    # Коэффициент перехода от qs к Rs по таблице 7.16 СП 24.13330-2011 изм.1
    t1 = 101.9716212978, 254.9290532445, 509.858106489, 764.7871597335, 1019.716212978, 1529.574319467, 2039.432425956, 3059.148638934
    t2 = 0.9, 0.8, 0.65, 0.55, 0.45, 0.35, 0.3, 0.2
    yy = qs
    if yy < 101.9716212978: yy = 101.9716212978
    B1 = interpoi(t1, t2, yy)
    #===============================================================================
    # Нормативное значение предельного сопротивления сваи на сжатие
    Funs = B1 * qs * asv + Funv
    # Расчетные коэффициенты
    ycs = 1.0; ycv = 0.8 # коэффициент условий работы сваи в грунте при сжатии/выдергивании
    ycg = 1.0 # коэффициент надежности по грунту
    # Несущую способность сваи при сжатии/выдергивании
    Fds = ycs * Funs / ycg
    Fdv = ycv * Funv / ycg
    #===============================================================================
    otv = int (dannie[8]) # Уровень ответственности
    if otv == 1: yn = 1.1 # коэффициент для повышенного уровня ответственности КС3
    if otv == 2: yn = 1.0 # коэффициент для нормального уровня ответственности КС2
    if otv == 3: yn = 1.0 # коэффициент для пониженного уровня ответственности КС1
    #===============================================================================
    if sum(torf) > 0:
        if ace > sum(ni_osadki): 
            ni_t = ni_tddd
        else: 
            ni_t = ni
    else:
        ni_t = ni
    Fyder = Funv                    # Расчетная удерживающая сила от морозного пучения
    ni = v_n[0]; Fneg = fu_n()      # Негативная сила трения оттаивающего грунта
    bi_tab = bii(fi); fi_tab = fi

    # Негативная сила трения осадки торфа
    if sum(torf) > 0:
        OTO = osadka_torfa(bi_tab, ni_osadki, fi, psi, torf, usv)
        if type(OTO) == tuple:
            Fosadki = OTO[0]
            bi_osadki = OTO[1]
            fi_osadki = OTO[2]
            ni_osadki = OTO[3]
            bfh_i = OTO[4]
            bfh = OTO[5]
            sum_bfh_i = OTO[6]
        else: Fosadki = 0
        if Fosadki > Fneg:
            Fnegra = Fosadki
        else: Fnegra = Fneg
    else: Fnegra = Fneg
    #===============================================================================
    maxs = (vsv + Ns + Fnegra)
    maxv = (Nv - vsv) if (Nv - vsv) >= 0 else 0.0
    Ncj = Fds / yn / 1.25   # Максимально допустимая нагрузка при сжатии
    Nvd = Fdv / yn / 1.6    # Максимально допустимая нагрузка при выдергивании
    #===============================================================================
    # Расчетная удельнуя касательная сила морозного пучения  по таблице Ж.1 СП 24.13330-2011 изм.1
    yy = ace
    if yy < 1.5: tfh = 11.0
    if yy > 3.0: tfh = 7.0
    if 1.5 <= yy <=3.0: t1 = 1.5, 2.5, 3.0; t2 = 11.0, 9.0, 7.0; tfh = interpoi(t1, t2, yy)
    #===============================================================================
    Apov = usv * ace; Fyderj = Fyder /1.1
    Npuch = kru * tfh * Apov - vsv*0.9
    #===============================================================================
    # ui.textEdit.clear() # Очиста textEdit
       
    def sopr_bok(bi_tab, ni_tab, fi_tab):
        bfh = []
        for  i in range(0, len(ni_tab)):
            if ni_tab[i] > 0:
                bfh_i = round((bi_tab[i] * fi_tab[i] * ni_tab[i]), 3)
                text_centr_black('слой № {}  -  {} ∙ {} т⁄м2 ∙ {} м = {} тс/м'.format(i+1, str(round(float(bi_tab[i]), 3)).ljust(5, '0'), str(round(float(fi_tab[i]), 3)).ljust(5, '0'), str(round(float(ni_tab[i]), 3)).ljust(4, '0'), str(float(bfh_i)).ljust(6, '0')))
                bfh.append(bfh_i)
            else: i += 1
        bfhj = [str(i) + ' тс/м' for i in bfh]
        bfhj = ' + '.join(bfhj)
        text_abzac('∑ ( βi ∙ fsi ∙ hi ) = {} = {} тс/м'.format(bfhj, round(sum(bfh),3)))
        return bfh    
    # --------------------------------------------------------------------------
    def sopr_bok_osadki(bi_osadki, fi_osadki, ni_osadki, bfh_i, bfh, sum_bfh_i):
        for i in range(0, len(ni_osadki)):
            if type(ni_osadki[i]) == list:
                for x in range(0, len(ni_osadki[i])):
                    text_centr_black('слой № {}  -  {} ∙ {} т⁄м2 ∙ {} м = {} тс/м'.format(i+1, str(round(float(bi_osadki[i]), 3)).ljust(5, '0'), str(round(float(fi_osadki[i][x]), 3)).ljust(5, '0'), str(round(float(ni_osadki[i][x]), 3)).ljust(4, '0'), str(float(bfh[i][x])).ljust(6, '0')))
            else:
                text_centr_black('слой № {}  -  {} ∙ {} т⁄м2 ∙ {} м = {} тс/м'.format(i+1, str(round(float(bi_osadki[i]), 3)).ljust(5, '0'), str(round(float(fi_osadki[i]), 3)).ljust(5, '0'), str(round(float(ni_osadki[i]), 3)).ljust(4, '0'), str(float(bfh[i])).ljust(6, '0')))
        bfhj = [str(i) + ' тс/м' for i in sum_bfh_i]; bfhj = ' + '.join(bfhj)
        text_abzac('∑ ( βi ∙ fsi ∙ hi ) = {} = {} тс/м'.format(bfhj, round(sum(sum_bfh_i), 3)))
        text_centr('Fneg = {} м ∙ {} тс/м = {} тс'.format(round(usv,3), round(sum(sum_bfh_i),3), round(Fosadki, 3)))

    # --------------------------------------------------------------------------
    # ui.textEdit.setFontWeight(100)
    text_centr('''Расчет одиночной сваи по результатам статического зондирования''')
    # ui.textEdit.setFontWeight(1)
    
    pix.save("image.png")
    # fff = os.getcwd() + '/image/image.png'
    # print(fff)
    # pix.save(fff)
    # ui.label_10.grab().save("image.png")
    paragraph = document.add_paragraph()
    paragraph.add_run('Геологический разрез', style='Intense Emphasis').bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_picture('image.png', width=Cm(5.0))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ui.textEdit.setFontUnderline(True)
    text_centr('Исходные данные:')
    # ui.textEdit.setFontUnderline(False)

    text_abzac_color('Параметры сваи:')
    text_abzac('\tДлина: {} м'.format(l1_ishod))
    text_abzac('\tОтметка верха: {} м'.format(l2_ishod))
    # if ui.radioButton.isChecked() == True:
    if True:
        text_abzac('\tДиаметр: {} м'.format(d1))
        text_abzac('\tТолщина стенки: {} м'.format(round(tsv,3)))
    else:
        text_abzac('\tСторона сваи: {} м'.format(d1))
        text_abzac('\tСобственный вес сваи: {} тс'.format(round(vsv,3)))

    text_abzac_color('Нагрузка на сваю:')
    text_abzac('\tСжимающая нагрузка на голову сваи N: {} тс'.format(round(Ns,3)))
    if maxv > 0.0: 
        text_abzac('\tВыдергивающая нагрузка на сваю N: {} тс'.format(round(Nv,3)))
    # if ui.radioButton.isChecked() == True:
    if True:
        text_abzac('\tКоэффициент надежности по нагрузке: для металла - {}, для бетона - {}'.format(knn_met, knn_beton))
        text_abzac('\tСобственный вес сваи ( металл ): {} ∙ {} тс = {} тс'.format(knn_met, round(msv,3), round(knn_met*msv,3)))
        text_abzac('\tСобственный вес сваи ( заполнение бетоном ): {} ∙ {} тс = {} тс'.format(knn_beton, round(bsv,3), round(knn_beton*bsv,3)))
        text_abzac('\tСобственный вес сваи: {} тс + {} тс = {} тс'.format(round(knn_met*msv,3), round(knn_beton*bsv,3), round(vsv,3)))
        text_abzac('\tРасчетная сжимающая нагрузка: Nс = {} тс + {} тс  = {} тс'.format(round(Ns,3), round(vsv,3), round(vsv + Ns,3)))
    else:
        text_abzac('\tКоэффициент надежности по нагрузки для бетона: {}'.format(knn_beton))
        text_abzac('\tСобственный вес сваи: {} тс ∙ {} тс = {} тс'.format(round(knn_beton,3), round(bsv,3), round(vsv,3)))
        text_abzac('\tРасчетная сжимающая нагрузка: Nс = {} тс + {} тс  = {} тс'.format(round(Ns,3), round(knn_beton*bsv,3), round(vsv + Ns, 3)))
    if maxv > 0.0:
        text_abzac('\tРасчетная выдергивающая нагрузка: Nв = {} тс - {} тс = {} тс'.format(round(Nv,3), round(vsv,3), round(Nv - vsv,3)))

    text_abzac_color('Параметры грунтов:')
    if Fneg > 0:
        text_abzac('\tПромерзаниe грунта - {} м'.format(ace))
    text_abzac('\tСреднее значение сопротивления грунта под наконечником зонда qs = {} т⁄м2'.format(round(qs,3)))
    for i in range(0, len(ni_tolchina)):
        # ui.textEdit.setFontWeight(100)
        text_abzac_color('Слой №{}:'.format(i+1))
        # ui.textEdit.setFontWeight(1)
        text_abzac('\tТолщина грунта hi = {} м'.format(round(ni_tolchina[i],3)))
        if fi[i] != 0:
            text_abzac('\tСопротивление грунта на боковой поверхности зонда fi = {} т⁄м2'.format(round(fi[i],3)))
        text_abzac('\tТип грунта - {}'.format(tigr[i]))
    # ui.textEdit.setFontUnderline(True)
    text_centr('Результаты расчета')
    # ui.textEdit.setFontUnderline(False)
    # ui.textEdit.append('')
    # ===========================================================
    if Fneg > 0:
        ni_tab = ni
        text_abzac('Ж.6    Расчет отрицательной силы трения оттаивающих грунтов на сваи')
        text_abzac('Отрицательную (негативную) силу трения оттаивающего грунта по боковой поверхности сваи можно определить по формуле:')
        text_centr('Fneg = u ∙ ∑ ( fni ∙ hni ) = u ∙ ∑ ( βi ∙ fsi ∙ hi )')
        text_abzac('где:')
        # if ui.radioButton.isChecked() == True:
        if True:
            text_abzac('u = 3.14 ∙ {} м = {} м - периметр поперечного сечения сваи;'.format(round(d1,4), round(usv,4)))
        else:
            text_abzac('u = {} м ∙ 4 = {} м - периметр поперечного сечения сваи;'.format(round(d1,4), round(usv,4)))
        text_abzac('hni = hi - толщина i-го слоя грунта, м;')
        text_abzac('fni = βi ∙ fsi - отрицательное трение i-го слоя оттаивающего грунта по боковой поверхности сваи;')
        text_abzac('βi - коэффициент, принимаемый по таблице 7.16 СП 24.13330.2011;')
        text_abzac('fsi - среднее сопротивление i-го слоя грунта на боковой поверхности зонда, тс/м2;')
        text_abzac('βi ∙ fsi ∙ hi :')
        sopr_b = sopr_bok(bi_tab, ni_tab, fi_tab)
        text_centr('Fneg = {} м ∙ {} тс/м = {} тс'.format(round(usv,3), round(sum(sopr_b),3), round(Fneg, 3)))
        # ui.textEdit.append('')
        
    if sum(torf) !=0:
        if Fosadki != 0:
            text_abzac('7.2.12    В пределах длины погруженной части сваи залегают напластования торфа толщиной более 30 см и планировка территории подсыпкой. Отрицательное трение осадки грунтов расположенных выше подошвы наинизшего (в пределах длины погруженной части сваи) слоя торфа, следует принимать:')
            text_centr('Fneg = u ∙ ∑ ( βi ∙ fsi ∙ hi )')
            text_abzac('где:')
            # if ui.radioButton.isChecked() == True:
            if True:
            
                text_abzac('u = 3.14 ∙ {} м = {} м - периметр поперечного сечения сваи;'.format(round(d1,4), round(usv,4)))
            else:
                text_abzac('u = {} м ∙ 4 = {} м - периметр поперечного сечения сваи;'.format(round(d1,4), round(usv,4)))
            text_abzac('hni = hi - толщина i-го слоя грунта, м;')
            text_abzac('fni = βi ∙ fsi - отрицательное трение i-го слоя оттаивающего грунта по боковой поверхности сваи;')
            text_abzac('βi - коэффициент, принимаемый по таблице 7.16, либо равный 1.0 для торфа, а также насыпи, fsi которой определено по таблице 7.3;')
            text_abzac('fsi - среднее сопротивление i-го слоя грунта на боковой поверхности зонда принятое с коэффициентом: для торфа - 1.0, для насыпи - 0.4, тс/м2;')
            text_abzac('βi ∙ fsi ∙ hi :')
            sopr_bok_osadki(bi_osadki, fi_osadki, ni_osadki, bfh_i, bfh, sum_bfh_i)
            # ui.textEdit.append('')
        # ===========================================================
    # if Fneg > 0:
    if Npuch > 0:
        ni_tab = ni_t
        text_abzac('Ж.4    Расчетное значение силы Frf, удерживающей сваи от выпучивания, определять по формуле:')
        text_centr('Frf = u ∑ ( fi ∙ hi ) = u ∙ ∑ ( βi ∙ fsi ∙ hi )')
        text_abzac('где:')
        # if ui.radioButton.isChecked() == True:
        if True:
                        
            text_abzac('u = 3.14 ∙ {} м = {} м - периметр поперечного сечения сваи.'.format(round(d1,4), round(usv,4)))
        else:
            text_abzac('u = {} м ∙ 4 = {} м - периметр поперечного сечения сваи.'.format(round(d1,4), round(usv,4)))
        text_abzac('βi - коэффициент, принимаемый по таблице 7.16 СП 24.13330.2011;')
        text_abzac('fi = fsi - расчетное сопротивление i-го слоя талого грунта сдвигу по поверхности сваи')
        text_abzac('hi - толщина i-го слоя талого грунта, расположенного ниже подошвы слоя промерзания-оттаивания, м;')
        text_abzac('βi ∙ fsi ∙ hi :') 
        sopr_b = sopr_bok(bi_tab, ni_tab, fi_tab)
        text_centr('Frf = {} ∙ {} = {} тс'.format(round(usv,3), round(sum(sopr_b),3), round(Fyder, 3)))
        # ui.textEdit.append('')

        text_abzac('Ж.2    Устойчивость свайных фундаментов на действие касательных сил морозного пучения грунтов надлежит проверять по условию:')
        text_centr('k ∙ τfh ∙ Afh - F ≤ Frf ∙ γc / γk')
        text_abzac('где:')
        text_abzac('τfh = {} тс/м2 - расчетная удельная касательная сила пучения, принимаемая согласно указаниям Ж.3 по таблице Ж.1 в зависимости от вида и характеристик грунта;'.format(round(tfh,3)))
        text_abzac('k = {} - коэффициент, корректирующий удельную касательную силу пучения, в зависимости от покрытия сваи в зоне промерзания грунтов;'.format(kru))
        # if ui.radioButton.isChecked() == True:
        if True:
            text_abzac('Afh = 3.14 ∙ {} м ∙ {} м = {} м2 - площадь боковой поверхности смерзания сваи в пределах расчетной глубины сезонного промерзания-оттаивания грунта или слоя искусственно замороженного грунта;'.format(round(d1,4), round(ace,3), round(Apov,3)))
        else:
            text_abzac('Afh = {} м ∙ 4 ∙ {} м = {} м2 - площадь боковой поверхности смерзания сваи в пределах расчетной глубины сезонного промерзания-оттаивания грунта или слоя искусственно замороженного грунта;'.format(round(d1,4), round(ace,3), round(Apov,3)))
        text_abzac('F = {} ∙ 0.9 = {} тс - расчетная нагрузка на сваю, принимаемая с коэффициентом 0,9 по наиболее невыгодному сочетанию нагрузок и воздействий, включая выдергивающие (ветровые, крановые и т. п.);'.format(round(vsv, 3), round(vsv*0.9, 3)))
        text_abzac('γc = 1.0 - коэффициент условий работы;')
        text_abzac('γk = 1.1 - коэффициент надежности;')
        text_abzac('Frf = {} тс - расчетное значение силы, удерживающей сваю от выпучивания вследствие трения его боковой поверхности о талый грунт, лежащий ниже расчетной глубины промерзания, кН, принять по указаниям Ж.4;'.format(round(Fyder, 3)))
        text_centr('{} ∙ {} тс/м2 ∙ {} м2 - {} тс ≤ {} тс ∙ 1.0 / 1.1'.format(kru, round(tfh,3), round(Apov,3), round(vsv*0.9,3), round(Fyder, 3)))
        text_centr('{} тс ≤ {} тс'.format(round(Npuch,3), round(Fyderj,3)))
        # ui.textEdit.append('')

    ni_tab = ni_t
    text_abzac('7.3.10    Частное значение предельного сопротивления забивной сваи в точке зондирования Fu, кН, следует определять по формуле:')
    if maxv > 0.0:
        text_abzac('для сжимающих нагрузок:')
        text_centr('fu = Rs ∙ A + f ∙ h ∙ u = β1 ∙ qs ∙ A + ∑ ( βi ∙ fsi ∙ hi ) ∙ u')
        # ui.textEdit.append('')
        text_abzac('для выдергивающих нагрузок:')
        text_centr('fu = f ∙ h ∙ u = ∙ A + ∑ ( βi ∙ fsi ∙ hi ) ∙ u')    
    else:
        text_centr('fu = Rs ∙ A + f ∙ h ∙ u = β1 ∙ qs ∙ A + ∑ ( βi ∙ fsi ∙ hi ) ∙ u')
    text_abzac('где:')
    text_abzac('Rs = β1 ∙ qs = {}  ∙ {} тc/м2 = {} тc/м2 - предельное сопротивление грунта под нижним концом сваи по данным зондирования в рассматриваемой точке;'.format(round(B1,4), round(qs,3), round(B1*qs,3)))
    text_abzac('β1 = {} - коэффициент перехода от qs к Rs, принимаемый по таблице 7.16 независимо от типа зонда по ГОСТ 19912;'.format(round(B1,4)))
    text_abzac('qs = {} тc/м2 - среднее значение сопротивления грунта под наконечником зонда;'.format(round(qs,3)))
    # if ui.radioButton.isChecked() == True:
    if True:
        text_abzac('А = 3.14 ∙ ( {} м )^2 ∙ 0.25 = {} м2 - площадь поперечного сечения натурной сваи;'.format(round(d1,4), round(asv,4))) # 3.14159265359 * d1*d1 * 0.25
    else:
        text_abzac('А = ( {} м )^2 = {} м2 - площадь поперечного сечения натурной сваи;'.format(round(d1,4), round(asv,4))) # d1*d1
    # if ui.radioButton.isChecked() == True:
    if True:
            text_abzac('u = 3.14 ∙ {} м = {} м - периметр поперечного сечения сваи;'.format(round(d1,4), round(usv,4)))
    else:
        text_abzac('u = {} м ∙ 4 = {} м - периметр поперечного сечения сваи;'.format(round(d1,4), round(usv,4)))
    text_abzac('h = {} м - глубина погружения сваи от поверхности грунта около сваи;'.format(round(sum(ni_tab), 3)))
    text_abzac('f = ∑ ( βi ∙ fsi ∙ hi ) / h - среднее значение предельного сопротивления грунта на боковой поверхности сваи по данным зондирования в рассматриваемой точке;')
    text_abzac('βi - коэффициент, принимаемый по таблице 7.16 СП 24.13330.2011;')
    text_abzac('fsi - среднее сопротивление i-го слоя грунта на боковой поверхности зонда, тс/м2 ;')
    text_abzac('hi - толщина i-го слоя грунта, м;')
    text_abzac('βi ∙ fsi ∙ hi :') 
    sopr_b = sopr_bok(bi_tab, ni_tab, fi_tab)
    if maxv > 0.0:
        text_abzac('для сжимающих нагрузок:')
        text_centr('fu = {} тc/м2 ∙ {} м2 + {} тc/м ∙ {} м = {} тс'.format(round(B1*qs,3), round(asv,4), round(sum(sopr_b),3), round(usv,3), round(Funs,3)))
        # ui.textEdit.append('')
        text_abzac('для выдергивающих нагрузок:')    
        text_centr('fu = {} тc/м ∙ {} м = {} тс'.format(round(sum(sopr_b),3), round(usv,3), round(Funv,3)))
    else:
        text_centr('fu = {} тc/м2 ∙ {} м2 + {} тc/м ∙ {} м = {} тс'.format(round(B1*qs,3), round(asv,4), round(sum(sopr_b),3), round(usv,3), round(Funs,3)))
    # ui.textEdit.append('')
    
    text_abzac('7.3.3    Несущую способность Fd, кН, свай по результатам их испытаний вдавливающей, выдергивающей и горизонтальной статическими нагрузками, а также по результатам их динамических испытаний следует определять по формуле:')
    text_centr('Fd = γc ∙ Fu,n / γc,g')
    text_abzac('где:')
    text_abzac('γc - коэффициент условий работы сваи ( в случае вдавливающих или горизонтальных нагрузок γc = 1.0; в случае выдергивающих нагрузок γc = 0.8 принимают по 7.2.5);')
    text_abzac('Fu,n = Fu = fu - нормативное значение предельного сопротивления сваи, определяемое в соответствии с 7.3.10;')
    text_abzac('γc,g = 1.0 - коэффициент надежности по грунту, принимаемый по указаниям 7.3.4.')
    if maxv > 0.0:
        text_abzac('для сжимающих нагрузок:')
        text_centr('Fd = {} ∙ {} тс / {} = {} тс'.format(1.0, round(Funs,3), 1.0, round(Fds,3)))
        # ui.textEdit.append('')
        text_abzac('для выдергивающих нагрузок:')
        text_centr('Fd = {} ∙ {} тс / {} = {} тс'.format(0.8, round(Funv,3), 1.0, round(Fdv,3)))
    else:
        text_centr('Fd = {} ∙ {} тс / {} = {} тс'.format(1.0, round(Funs,3), 1.0, round(Fds,3)))
    # ui.textEdit.append('')
    
    text_abzac('7.1.11    Сваю в составе фундамента и одиночную по несущей способности грунта основания следует рассчитывать исходя из условия:')
    text_centr('N ≤ Fd / γn / γc,g')
    text_abzac('где:')
    if maxv > 0.0:
        text_abzac('N - расчетная нагрузка, передаваемая на сваю от наиболее невыгодного сочетания нагрузок, действующих на фундамент, определяемая в соответствии с 7.1.12;')
        text_abzac('для сжимающих нагрузок:')
        if Fnegra > 0: text_abzac('\tN = Nc + Fneg = {} тс + {} тс = {} тс'.format(round(vsv + Ns,3), round(Fnegra,3), round(maxs, 3)))
        else: text_abzac('\tN = Nc = {} тс'.format(round(maxs, 3)))
        text_abzac('для выдергивающих нагрузок:')
        text_abzac('\tN = Nв = {} тс;'.format(round(maxv,3)))
    else:
        if Fnegra > 0: text_abzac('N = Nc + Fneg = {} тс + {} тс = {} тс - расчетная нагрузка, передаваемая на сваю от наиболее невыгодного сочетания нагрузок, действующих на фундамент, определяемая в соответствии с 7.1.12;'.format(round(vsv + Ns,3), round(Fnegra,3), round(maxs, 3)))
        else: text_abzac('N = Nc = {} тс - расчетная нагрузка, передаваемая на сваю от наиболее невыгодного сочетания нагрузок, действующих на фундамент, определяемая в соответствии с 7.1.12;'.format(round(maxs, 3)))
    text_abzac('Fd - несущая способность сваи, определяемая в соответствии с 7.3.3;')
    text_abzac('γn = {} - коэффициент надежности по ответственности сооружения, принимаемый по ГОСТ 27751, но не менее 1.0;'.format(yn))
    text_abzac('γc,g = 1,25 / 1,6 - коэффициент надежности по грунту при сжимающей / выдергивающей нагрузке.')
    if maxv > 0.0:
        text_abzac('для сжимающих нагрузок:')
        text_centr('{} тс ≤ {} тс / {} / 1.25 = {} тс '.format(round(maxs,3), round(Fds,3), yn, round(Ncj,3)))
        # ui.textEdit.append('')
        text_abzac('для выдергивающих нагрузок:')
        text_centr('{} тс ≤ {} тс / {} / 1.6 = {} тс '.format(round(maxv,3), round(Fdv,3), yn, round(Nvd,3)))
    else:
        text_centr('{} тс ≤ {} тс / {} / 1.25 = {} тс '.format(round(maxs,3), round(Fds,3), yn, round(Ncj,3)))
    # ===================================================================================================================
    text_abzac('Вывод:')
    if Fnegra > 0:
        text_abzac('Расчетная сжимающая нагрузка на сваю с учетом негативной силы трения грунтов: {} тс'.format(round(maxs, 2)))
    else: 
        text_abzac('Расчетная сжимающая нагрузка на сваю: {} тс'.format(round(maxs, 2)))

    if maxs > Ncj:
        # ui.textEdit.setTextColor(QtGui.QColor (255, 0, 0))
        text_abzac('Максимально допустимая нагрузка при сжатии: ' + str(round (Ncj, 2)) + ' тс')
        # ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
    else: 
        text_abzac('Максимально допустимая нагрузка при сжатии: ' + str(round (Ncj, 2)) + ' тс')
    
    if maxv > 0.0:
        text_abzac('Расчетная выдергивающая нагрузка на сваю: {} тс'.format(round(maxv, 2))) 
        if maxv > Nvd:
            # ui.textEdit.setTextColor(QtGui.QColor (255, 0, 0))
            text_abzac('Максимально допустимая нагрузка при выдергивании: {} тс'.format(round(Nvd, 2))) 
            # ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
        else:
            text_abzac('Максимально допустимая нагрузка при выдергивании: {} тс'.format(round(Nvd, 2))) 
    
    if Npuch > 0:
    # if Fneg > 0:
        text_abzac('Расчетная выдергивающая нагрузка от сил морозного пучения грунта: {} тс'.format(round(Npuch, 2)))
        if Npuch > Fyderj:
            # ui.textEdit.setTextColor(QtGui.QColor (255, 0, 0))
            text_abzac('Расчетная удерживающая сила по боковой поверхности свай: {} тс'.format(round(Fyderj, 2)))
            # ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
        else:
            text_abzac('Расчетная удерживающая сила по боковой поверхности свай: {} тс'.format(round(Fyderj, 2)))
    
    
    
    
    
    # print(document)
    
    '''# Используя следующий код, можно получить весь текст документа:'''
    # text = []
    # for paragraph in document.paragraphs:
    #     text.append(paragraph.text)
    # print('\n\n'.join(text))
    
    # text = []
    # with open("result.html", "w", encoding='utf-8') as file:
    #     for paragraph in document.paragraphs:
    #         file.write(paragraph.text + '\n')
    

    document.save('result.docx')

    with open('document.html', 'w', encoding='utf-8') as file:
        file.write(textOtchet)
        file.close()    
    # return document
    # print(textOtchet)
    print('----------------------------------')
    print('----------------------------------')
    
    return textOtchet
# # ----------------------------------------------------------------------------------------
# def text_abzac_color_000(x):
#     ui.textEdit.setTextColor(QtGui.QColor (0, 100, 150))
#     ui.textEdit.setFontWeight(100) # жирный текст
#     ui.textEdit.append('{}'.format(x))
#     ui.textEdit.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
#     ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
#     ui.textEdit.setFontWeight(1) # жирный текст
#     # ui.textEdit.append('')

# def text_abzac_color_111(x):
#     ui.textEdit.setTextColor(QtGui.QColor (0, 100, 150))
#     ui.textEdit.setFontWeight(100) # жирный текст
#     ui.textEdit.insertPlainText('{}'.format(x))
#     ui.textEdit.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
#     ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
#     ui.textEdit.setFontWeight(1) # жирный текст

# def sreznull(x):
#     while x[-1] == 0: del x[-1]

# def beznull(x):
#     for i in range(0, len(x)-1):
#         for z in range(0, len(x[i])):
#             if x[i][z] == 0:
#                 x[i][z] = ''
#     return x

# myDocumentsPath = QtCore.QDir.homePath() + "\Documents"
# savePathFile = os.getcwd() + "\savePath.file"

# """Если файл с сохраненным путем существует и не пустой"""
# if os.path.exists(savePathFile) == True and savePathFile != '':
#     f = open(savePathFile, 'rb')        # чтение из файла
#     loadPath = pickle.load(f)           # извлекаем ообъект из файла
#     f.close()
# else: 
#     loadPath = myDocumentsPath

# def writeFail(katalog):
#     global loadPath
#     '''Сохраняем путь к последнему сохараненному файлу в корневай папки программы'''
#     f = open(savePathFile, 'wb')                # Запись в файл
#     pickle.dump(katalog, f)                     # помещаем объект в файл
#     f.close()

# def soxranka():
#     global loadPath
#     ff = QtWidgets.QFileDialog.getSaveFileName(Form, 'Сохранить как', loadPath, 'Statzond ( *.stz )')
#     katalog = ff[0]

#     if '-stz' not in katalog:
#         katalog = katalog[:-4] + '-stz' + katalog[-4:]

#     if katalog != '':
#         failName = os.path.split(katalog)[1]
#         loadPath = writeFail(katalog)
        
#         sbor_dannih()
#         savex = [ige_xap, qi_xap, fi_xap, dannie, ige_skv, nni, sechenie]
#         f = open(katalog, 'wb') # Запись в файл
#         pickle.dump(savex, f)   # помещаем объект в файл
#         f.close()
#         del savex               # уничтожаем переменную savex  
#         # ------------------------------------------------------------
#         section = document.sections[0]
#         section.top_margin  = Cm(2)     # верхнее поле
#         section.bottom_margin = Cm(2)   # нижнее поле
#         section.left_margin = Cm(3)     # левое поле
#         section.right_margin = Cm(1.5)    # правое поле
#         # ------------------------------------------------------------
#         try:
#             document.save('{}.docx'.format(katalog[:-4]))
#         except:
#             # error_show(f"Файл {katalog} открыт. Закройте файл для перезаписи . . .")
#             print(f"Файл {katalog} открыт. Закройте файл для перезаписи . . .")
#             return
#         Form.setWindowTitle(_translate("Form", "Pail - {}".format(failName)))
#         ui.label_3.setText(_translate("Form", "Сохранение"))
#         ui.textEdit.setText('')
#         ui.textEdit.insertPlainText('Данные сохранены в файл: ')
#         text_abzac_color_111(f'{failName}')
#         ui.textEdit.append('')
#         ui.textEdit.insertPlainText('Нажмите кнопку - ')
#         text_abzac_color_111(f'Расчет')
#     else: return None

# def otkrivalka():
#     global loadPath
#     ff = QtWidgets.QFileDialog.getOpenFileName(Form, 'Открыть', loadPath, 'Statzond ( *.stz )')
#     katalog = ff[0]
    
#     if katalog != '':
#         failName = os.path.split(katalog)[1]
#         loadPath = writeFail(katalog)

#         f = open(katalog, 'rb') # чтение из файла
#         loadx = pickle.load(f) # извлекаем ообъект из файла
#         for i in range(0, len(loadx)-1):
#             for z in range(0, len(loadx[i])):
#                 if loadx[i][z] == 0:
#                     loadx[i][z] = ''
#         def vvod_open(a, b, c):
#             _translate = QtCore.QCoreApplication.translate
#             xxx = eval('[\'\' for z in range(0, ui.tableWidget{}.rowCount())]'.format(a))
#             for i in range(0, len(xxx)):
#                 eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", str({})))'.format(a, i, b, xxx[i]))
#             for i in range(0, len(c)):
#                 # eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", str({})))'.format(a, i, b, c[i]))
#                 eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", "{}"))'.format(a, i, b, c[i]))

#         vvod_open('', 0, loadx[0])
#         vvod_open('', 1, loadx[1])
#         vvod_open('', 2, loadx[2])
#         vvod_open('_2', 0, loadx[3])
#         vvod_open('_4', 0, loadx[4])
#         vvod_open('_4', 1, loadx[5])
#         ui.radioButton.setChecked(loadx[6])
#         del loadx # уничтожаем переменную savex  
#         ui.label_3.setText(_translate("Form", "Открытие"))
#         ui.textEdit.setText('')
#         ui.textEdit.insertPlainText('Данные закгужены из файла: ')
#         text_abzac_color_111(f'{failName}')
#         Form.setWindowTitle(_translate("Form", "Pail - {}".format(failName)))
#         ui.textEdit.append('')
#         ui.textEdit.insertPlainText('Нажмите кнопку - ')
#         text_abzac_color_111(f'Расчет')
#         return None
#     else: return None
# # ----------------------------------------------------------------------------------------
# def text_centr_000(x):
#     ui.textEdit.append('')
#     ui.textEdit.setTextColor(QtGui.QColor (0, 100, 150)) # цвет текста
#     ui.textEdit.setFontWeight(100) # жирный текст
#     ui.textEdit.append(x)
#     ui.textEdit.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
#     ui.textEdit.setFontWeight(1) # убираем жирный текст
#     ui.textEdit.setTextColor(QtGui.QColor (0, 0, 0))
# def text_abzac_000(x):
#     ui.textEdit.append('       {}'.format(x))
#     ui.textEdit.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева

# def spravka():
#     ui.label_3.setText(_translate("Form", "Справка"))
#     ui.textEdit.setText('')
#     text_abzac_color_000('       ИГЭ')
#     ui.textEdit.insertPlainText(' – инженерно-геологический элемент')
#     text_abzac_color_000('       qsi')
#     ui.textEdit.insertPlainText(' – среднее значение сопротивления i-го слоя грунта под наконечником зонда')
#     text_abzac_color_000('       fi')
#     ui.textEdit.insertPlainText(' – среднее сопротивление i-го слоя грунта на боковой поверхности зонда')
#     text_abzac_color_000('       H')
#     ui.textEdit.insertPlainText(' – толщина i-го слоя грунта')

#     text_abzac_000('Поскольку сопротивление под конусом зонда в глинистых грунтах больше чем у торфа, но меньше чем у песчаных грунтов, принимаем следующие граничные параметры:')
#     text_abzac_color_000('       qsi max торф')
#     ui.textEdit.insertPlainText(' – максимальное сопротивление торфа, принимаем меньше минимального значения глинистых грунтов;')
#     text_abzac_color_000('       qsi min песок')
#     ui.textEdit.insertPlainText(' – минимальное сопротивление песка, принимаем больше максимального значения глинистых грунтов.')
#     text_abzac_000('Для корректного определения типа грунта в таблице " Характеристики грунтов " в колонке " qsi " принимаем для песков значение больше или равно " qsi min песок ", а для торфа значение меньше или равно " qsi max торф ".')
#     text_abzac_000('Для выделения песчаных грунтов в насыпные в таблице " Характеристики грунтов " в колонке " fi " указываем значение "0" (ноль).')
#     ui.textEdit.append('')
#     text_abzac_000('При сохранении Расчета сохраняются два файла, один из них в формате *.docx')
#     # text_abzac_000('P.S.: Вы можете выразить благодарность автору программы оценив его старания по указанной ссылке: https://money.yandex.ru/to/4100115068617256')
#     # ui.textEdit.append('')
#     # ui.textEdit.append('')
#     # ui.textEdit.textCursor().insertImage('QRCodY.png')
#     # ui.textEdit.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    
#     # text_abzac_000('')

# def QRCods():
#     import numpy
#     import png
#     import QRCod
#     fff = png.from_array(QRCod.a, mode="L")
#     fff.save('QRCodY.png') # works
#     # ui.label_10.setPixmap(QtGui.QPixmap("QRCodY.png"))
# QRCods()
# ----------------------------------------------------------------------------------------
# '''Авторизация'''
# noyblock = 'no'
# def block():
#     global noyblock
#     import datetime
#     dt_now = datetime.datetime.now().day
#     passeord = QtWidgets.QInputDialog.getText(Form, 'Авторизация', 'Введите пароль')
#     if str(dt_now) in passeord[0]:
#         noyblock = 'Yes'
#     else:
#         noyblock = 'no'
#         QtWidgets.QMessageBox.information(Form, 'Авторизация', 'Неверный пароль')

# def raschet_block():
#     global noyblock
#     if noyblock == 'no':
#         block()
#     if noyblock == 'Yes':
#         raschet()
# ----------------------------------------------------------------------------------------
# ----------------------------------------------------------------------------------------
# def clear_tab():
#     xxx = ['' for z in range(0, ui.tableWidget_4.rowCount())]
#     for s in range(ui.tableWidget_4.columnCount()):
#         for i in range(0, len(xxx)):
#             eval (f'ui.tableWidget_4.item({i}, {s}).setText(_translate("Form", str({xxx[i]})))')

# def loadGeo():
#     try:
#         import ReadTab
#         TabData = ReadTab.Read()[3:]
#         ige = [x[0] for x in TabData]
#         qqq = [x[2] if x[2] != '-' else '' for x in TabData]
#         fff = [x[3] if x[3] != '-' else '' for x in TabData]
#         for i in range(0, len(ige)):
#             eval (f'ui.tableWidget.item({i}, {0}).setText(_translate("Form", "{ige[i]}"))')
#             eval (f'ui.tableWidget.item({i}, {1}).setText(_translate("Form", "{qqq[i]}"))')
#             eval (f'ui.tableWidget.item({i}, {2}).setText(_translate("Form", "{fff[i]}"))')
#     except: pass

# def writeGeo():
#     try:
#         import WriteTab
#         WriteTab.write()
#     except: pass

'''Растягиваем все колонки в таблице симметрично'''
# def viravcolumntabAll(tab):
#     header = eval('ui.{}.horizontalHeader()'.format(tab))
#     columnCount = eval('ui.{}.columnCount()'.format(tab))
#     for i in range(columnCount):
#         header.setSectionResizeMode(i, QtWidgets.QHeaderView.Stretch)
# viravcolumntabAll('tableWidget')
# viravcolumntabAll('tableWidget_2')
# viravcolumntabAll('tableWidget_4')
# ----------------------------------------------------------------------------------------
# ----------------------------------------------------------------------------------------
# ui.pushButton_11.clicked.connect (writeGeo)
# ui.pushButton_9.clicked.connect (loadGeo)
# ui.pushButton_8.clicked.connect (clear_tab)
# # ui.pushButton.clicked.connect (raschet_block)
# ui.pushButton.clicked.connect (raschet)
# ui.pushButton_1.clicked.connect (soxranka)
# ui.pushButton_2.clicked.connect (otkrivalka)
# ui.pushButton_3.clicked.connect (spravka)
# ----------------------------------------------------------------------------------------
# if __name__ == "__main__":
#     sys.exit(app.exec_())
# ----------------------------------------------------------------------------------------

