# import os
# import base64
# from docx import Document
# import docx
# from lxml import etree


# # Открываем документ
# doc = Document('example.docx')

# # Создаем корневой элемент html
# html = etree.Element('html')

# # Создаем элемент head с заголовком
# head = etree.SubElement(html, 'head')
# title = etree.SubElement(head, 'title')
# title.text = 'Example Document'

# # Создаем элемент body и добавляем его в корневой элемент
# body = etree.SubElement(html, 'body')
# html.append(body)

# # Цикл по всем элементам в документе
# for element in doc.element.body:
#     print(element.text)
#     # Если элемент является параграфом
#     if isinstance(element.text, docx.text.paragraph.Paragraph):
#         # Создаем элемент p и добавляем его в элемент body
#         p = etree.SubElement(body, 'p')
#         # Добавляем текст параграфа в элемент p
#         p.text = element.text
        
#     # Если элемент является таблицей
#     elif isinstance(element, docx.table._Table):
#         # Создаем элемент table и добавляем его в элемент body
#         html_table = etree.SubElement(body, 'table')
#         # Цикл по всем строкам в таблице
#         for row in element.rows:
#             # Создаем элемент tr и добавляем его в элемент table
#             html_row = etree.SubElement(html_table, 'tr')
#             # Цикл по всем ячейкам в строке
#             for cell in row.cells:
#                 # Создаем элемент td и добавляем его в элемент tr
#                 html_cell = etree.SubElement(html_row, 'td')
#                 # Добавляем текст ячейки в элемент td
#                 html_cell.text = cell.text
#     # Если элемент является списком
#     elif isinstance(element, docx.text.list.List):
#         # Создаем элемент ul и добавляем его в элемент body
#         ul = etree.SubElement(body, 'ul')
#         # Цикл по всем элементам в списке
#         for item in element.items:
#             # Создаем элемент li и добавляем его в элемент ul
#             li = etree.SubElement(ul, 'li')
#             # Добавляем текст элемента в элемент li
#             li.text = item.text
#     # Если элемент является изображением
#     elif isinstance(element, docx.shape.InlineShape):
#         # Создаем элемент img и добавляем его в элемент body
#         img = etree.SubElement(body, 'img')
#         # Получаем бинарные данные изображения
#         img_data = element._blob
#         # Устанавливаем атрибут src элемента img
#         img.set('src', 'data:image/png;base64,' + base64.b64encode(img_data).decode('utf-8'))

# # Создаем объект ElementTree и записываем его в файл example.html
# tree = etree.ElementTree(html)
# tree.write('example.html', pretty_print=True, xml_declaration=True, encoding='UTF-8')



# import mammoth

# with open('document.docx', 'rb') as docx_file:
#     result = mammoth.convert_to_html(docx_file)
#     html = result.value # получить HTML-код
    
# with open('document.html', 'w', encoding='UTF-8') as html_file:
#     html_file.write(html)
# # отправить html по почте или вставить на веб-страницу




# import docx

# import os
# import base64
# from docx import Document
# from lxml import etree


# # Открываем документ
# doc = Document('example.docx')

# # Создаем корневой элемент html
# html = etree.Element('html')

# # Создаем элемент head с заголовком
# head = etree.SubElement(html, 'head')
# title = etree.SubElement(head, 'title')
# title.text = 'Example Document'

# # Создаем элемент body и добавляем его в корневой элемент
# body = etree.SubElement(html, 'body')
# html.append(body)

# # Цикл по всем элементам в документе
# for element in doc.element.body:
#     # Если элемент является параграфом
#     if isinstance(element, docx.text.paragraph.Paragraph):
#         # Создаем элемент p и добавляем его в элемент body
#         p = etree.SubElement(body, 'p')
#         # Добавляем текст параграфа в элемент p
#         p.text = element.text

#     # # Если элемент является изображением
#     # elif isinstance(element, docx.shape.InlineShape):
#     #     # Создаем элемент img и добавляем его в элемент body
#     #     img = etree.SubElement(body, 'img')
#     #     # Получаем бинарные данные изображения
#     #     img_data = element._inline.graphic.graphicData.pic.blipFill.blip.blob
#     #     # Устанавливаем атрибут src элемента img
#     #     img.set('src', 'data:image/png;base64,' + base64.b64encode(img_data).decode('utf-8'))

# # Создаем объект ElementTree и записываем его в файл example.html
# tree = etree.ElementTree(html)
# tree.write('example.html', pretty_print=True, xml_declaration=True, encoding='UTF-8')


# import os
# import base64
# from docx import Document
# from lxml import etree

# # Открываем документ
# doc = Document('example.docx')

# # Создаем корневой элемент html
# html = etree.Element('html')

# # Создаем элемент head с заголовком
# head = etree.SubElement(html, 'head')
# title = etree.SubElement(head, 'title')
# title.text = 'Example Document'

# # Создаем элемент body и добавляем его в корневой элемент
# body = etree.SubElement(html, 'body')
# html.append(body)



# # Создаем элемент style и добавляем его в элемент head
# style = etree.SubElement(head, 'style')
# style.set('type', 'text/css')
# head.append(style)

    
# # Цикл по всем параграфам в документе
# for para in doc.paragraphs:
#     p = etree.SubElement(body, 'p')
#     p.text = para.text
    
#     style = para.style
#     css = ''
    


    
#     if style.font.name is not None:
#         css += 'font-family:' + style.font.name + ';'
#     if style.font.size is not None:
#         css += 'font-size:' + str(style.font.size.pt) + 'pt;'
#     if style.font.color.rgb is not None:
#         css += 'color:#' + style.font.color.rgb[2:] + ';'
#     if style.font.bold is not None:
#         css += 'font-weight:bold;'
#     if style.font.italic is not None:
#         css += 'font-style:italic;'
#     if style.font.underline is not None:
#         css += 'text-decoration:underline;'
#     if style.paragraph_format.alignment is not None:
#         css += 'text-align:' + style.paragraph_format.alignment + ';'
        
#     if css != '':
#         p.set('style', css)

# # Создаем объект ElementTree и записываем его в файл example.html
# tree = etree.ElementTree(html)
# tree.write('example.html', pretty_print=True, xml_declaration=True, encoding='UTF-8')




# from docx import Document

# def get_css_from_style(style):
#     css = ''
#     if style.font.name is not None:
#         css += 'font-family:' + style.font.name + ';'
#     if style.font.size is not None:
#         css += 'font-size:' + str(style.font.size.pt) + 'pt;'
#     if style.font.color.rgb is not None:
#         css += 'color:#' + style.font.color.rgb[2:] + ';'
#     if style.font.bold is not None:
#         css += 'font-weight:bold;'
#     if style.font.italic is not None:
#         css += 'font-style:italic;'
#     if style.font.underline is not None:
#         css += 'text-decoration:underline;'
#     if style.paragraph_format.alignment is not None:
#         css += 'text-align:' + style.paragraph_format.alignment + ';'
#     return css

# doc = Document('example.docx')

# css_list = []

# for style in doc.styles:
#     print(style)
#     print(style.font.size)
#     if style.name != 'Normal':
#         css = '.' + style.name + '{' + get_css_from_style(style) + '}'
#         css_list.append(css)

# css_string = '\n'.join(css_list)

# with open('example.css', 'w') as f:
#     f.write(css_string)


import os
import base64
from docx import Document
from lxml import etree
from docx.enum.style import WD_STYLE_TYPE

# Открываем документ
doc = Document('example.docx')

# Создаем корневой элемент html
html = etree.Element('html')

# Создаем элемент head с заголовком
head = etree.SubElement(html, 'head')
title = etree.SubElement(head, 'title')
title.text = 'Example Document'

# Создаем элемент body и добавляем его в корневой элемент
body = etree.SubElement(html, 'body')
html.append(body)


for para in doc.paragraphs:
    font_name = para.style.font.bold
    print(f'Font name: {font_name}')
    
# Цикл по всем параграфам в документе
# for para in doc.paragraphs:
#     p = etree.SubElement(body, 'p')
#     p.text = para.text
    
#     style = para.style
    
    
    
    
    
    # all_styles = p.style
    # paragraph_styles = [s for s in all_styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    # for style in paragraph_styles:
    #     print(style.name)


    # print(f'style = {style}')
    # css = ''
    # print(f'style.name = {style.name}')
    # style = doc.styles[style.name]
    # print(f'style.name = {style.name}')
    # print(f'style.size = {style.size}')
    
    # if style.font.name:
    #     css += 'font-family:' + style.font.name + ';'
    # if style.font.size:
    #     css += 'font-size:' + str(style.font.size.pt) + 'pt;'
    # if style.font.color.rgb:
    #     css += 'color:#' + style.font.color.rgb[2:] + ';'
    # if style.font.bold:
    #     css += 'font-weight:bold;'
    # if style.font.italic:
    #     css += 'font-style:italic;'
    # if style.font.underline:
    #     css += 'text-decoration:underline;'
    # if style.paragraph_format.alignment:
    #     css += 'text-align:' + style.paragraph_format.alignment + ';'
        
    # if css != '':
    #     p.set('style', css)

# Создаем объект ElementTree и записываем его в файл example.html
tree = etree.ElementTree(html)
tree.write('example.html', pretty_print=True, xml_declaration=True, encoding='UTF-8')


